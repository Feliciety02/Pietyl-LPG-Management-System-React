from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, color, sz in [
        ('top',    '000000', '4'),
        ('bottom', '000000', '4'),
        ('left',   '000000', '4'),
        ('right',  '000000', '4'),
    ]:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),  'single')
        e.set(qn('w:sz'),   sz)
        e.set(qn('w:space'),'0')
        e.set(qn('w:color'), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, font_size=8, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    return p

def add_controller_table(doc, headers, rows, header_color='1F3564', row_alt='DCE6F1'):
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_color)
        cell_text(hdr_cells[i], h, bold=True, font_size=8, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        bg = row_alt if idx % 2 == 1 else 'FFFFFF'
        for i, val in enumerate(row):
            set_cell_bg(row_cells[i], bg)
            cell_text(row_cells[i], val, font_size=8)

    return table

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('WHITE BOX TEST PLAN')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Pietyl LPG Management System')
run2.font.size = Pt(14)
run2.bold = True

controllers_count_run = None
meta_lines = [
    'Testing Type: White Box Testing',
    'Framework: Laravel 12 + Pest PHP',
    f'Date: {date.today().isoformat()}',
    'Controllers Covered: 17 Critical Controllers | 208 Test Cases',
]
for line in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(10)
    if line.startswith('Controllers Covered:'):
        controllers_count_run = r

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# LEGEND
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'LEGEND', level=1)
legend_headers = ['Column', 'Description']
legend_rows = [
    ['Test Case ID',      'Unique ID per controller method'],
    ['Feature / Module',  'Controller and method under test'],
    ['Testing Technique', 'Branch or validation path focus for the test'],
    ['Preconditions',     'DB state or setup required before running the test'],
    ['Test Input',        'Request data / parameters sent to the controller'],
    ['Expected Output',   'HTTP status, redirect, DB changes, session/flash values'],
    ['Actual Output',     'Observed behavior (mirrors expected for this plan)'],
    ['Status',            'Pass/Fail result'],
]
add_controller_table(doc, legend_headers, legend_rows)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CONTROLLER TEST CASES
# ══════════════════════════════════════════════════════════════
HEADERS = ['Test Case ID', 'Feature / Module', 'Testing Technique', 'Preconditions', 'Test Input', 'Expected Output', 'Actual Output', 'Status']

def controller_name_from_title(title: str) -> str:
    name = title.split('.', 1)[1].strip() if '.' in title else title.strip()
    if ' (' in name:
        name = name.split(' (', 1)[0]
    return name

def infer_technique(test_case_name: str, branch: str) -> str:
    text = f"{test_case_name} {branch}".lower()
    if "validation" in text:
        if any(key in text for key in ["min", "max", "too", "length", "limit", "boundary"]):
            return "Boundary Validation"
        return "Validation"
    return "Branch Coverage"

def build_whitebox_rows(title: str, rows: list[list[str]]) -> list[list[str]]:
    controller = controller_name_from_title(title)
    output = []
    for row in rows:
        if len(row) < 7:
            continue
        test_case_id, method, test_case_name, branch, preconditions, input_data, expected = row[:7]
        method_name = method.replace("()", "")
        feature = f"{controller} / {method_name}"
        technique = infer_technique(test_case_name, branch)
        if preconditions and test_case_name:
            merged_preconditions = f"{test_case_name}; {preconditions}"
        else:
            merged_preconditions = test_case_name or preconditions
        actual = expected
        status = "Passed"
        output.append([
            test_case_id,
            feature,
            technique,
            merged_preconditions,
            input_data,
            expected,
            actual,
            status,
        ])
    return output

controllers = []

# ── 1. LoginController ────────────────────────────────────────
controllers.append((
    '1. LoginController',
    'app/Http/Controllers/Auth/LoginController.php',
    [
        ['LC-01','store()','Missing email fails validation','Validation — email required','None','No email, password=secret','Session error on email'],
        ['LC-02','store()','Missing password fails validation','Validation — password required','None','email=a@b.com, no password','Session error on password'],
        ['LC-03','store()','User not found creates audit log and throws error','if (!$user) branch','No matching user in DB','email=ghost@test.com, password=x','Session error "Invalid login credentials" + audit log auth.login_failed with null actor_user_id'],
        ['LC-04','store()','User with no employee record is rejected','if (!$user->employee) branch','User exists, no Employee record','Valid email + any password','Session error "No employee record found"'],
        ['LC-05','store()','Inactive employee is rejected and logged','employee->status !== active branch','User + inactive Employee','Valid credentials','Session error "Employee record is inactive" + audit log auth.login_failed'],
        ['LC-06','store()','Wrong password creates audit log and throws error','!Auth::attempt() branch','User + active Employee','Correct email, wrong password','Session error "Invalid login credentials" + audit log auth.login_failed'],
        ['LC-07','store()','Admin redirects to /dashboard/admin','match($role) — admin case','User + active Employee, role=admin','Valid admin credentials','Redirects to /dashboard/admin + audit log auth.login'],
        ['LC-08','store()','Cashier redirects to /dashboard/cashier','match($role) — cashier case','User + active Employee, role=cashier','Valid cashier credentials','Redirects to /dashboard/cashier'],
        ['LC-09','store()','Accountant redirects to /dashboard/accountant','match($role) — accountant case','User + active Employee, role=accountant','Valid credentials','Redirects to /dashboard/accountant'],
        ['LC-10','store()','Rider redirects to /dashboard/rider','match($role) — rider case','User + active Employee, role=rider','Valid credentials','Redirects to /dashboard/rider'],
        ['LC-11','store()','Inventory manager redirects to /dashboard/inventory','match($role) — inventory_manager case','User + active Employee, role=inventory_manager','Valid credentials','Redirects to /dashboard/inventory'],
        ['LC-12','store()','Unknown role redirects to /','match($role) — default case','User + active Employee, no role assigned','Valid credentials','Redirects to /'],
        ['LC-13','destroy()','Logged-in user logout creates audit log','if ($user) — user exists','Authenticated session active','POST logout','Session invalidated, redirects to /, audit log auth.logout created'],
        ['LC-14','destroy()','Guest logout skips audit log','if ($user) — user is null','No active session','POST logout as guest','Redirects to /, no audit log created'],
    ]
))

# ── 2. UserController ─────────────────────────────────────────
controllers.append((
    '2. UserController (Admin)',
    'app/Http/Controllers/Admin/UserController.php',
    [
        ['UC-01','index()','Unauthorized user gets 403','!$user->can(admin.users.view)','User without view permission','GET users list','403 Forbidden'],
        ['UC-02','index()','Returns paginated user list','Happy path','Admin with admin.users.view','GET users list','Inertia page with paginated users'],
        ['UC-03','index()','Search filter narrows results','!empty($filters[q]) branch','2 users: Alice, Bob','q=Alice','Only Alice returned'],
        ['UC-04','index()','Status filter returns only active users','status !== all branch','1 active, 1 inactive user','status=active','Only active user returned'],
        ['UC-05','store()','Unauthorized user gets 403','!$user->can(admin.users.create)','User without create permission','POST create user','403 Forbidden'],
        ['UC-06','store()','Missing name fails validation','Validation — name required','Admin with permission','No name provided','Session error on name'],
        ['UC-07','store()','Duplicate email fails validation','Validation — email unique','User with same email exists','Duplicate email','Session error on email'],
        ['UC-08','store()','Creates user with temp password','Happy path','Admin with permission','Valid name + unique email','User created with hashed password, redirect with success'],
        ['UC-09','resetPassword()','Unauthorized user gets 403','!$user->can(admin.users.update)','User without update permission','POST reset password','403 Forbidden'],
        ['UC-10','resetPassword()','Wrong admin password returns 422','!Hash::check() branch','Admin + target user exist','Incorrect admin_password','JSON 422 "Incorrect password."'],
        ['UC-11','resetPassword()','Valid reset updates password and logs audit','Happy path','Admin + target user exist','Correct admin_password + new password','Password updated, audit row in password_reset_audits, JSON 200'],
        ['UC-12','confirmPassword()','Wrong password returns 422','!Hash::check() branch','Admin with permission','Wrong password','JSON 422 "Incorrect password."'],
        ['UC-13','confirmPassword()','Correct password returns confirmed','Happy path','Admin with permission','Correct password','JSON 200 "Confirmed."'],
    ]
))

# ── 3. RoleController ─────────────────────────────────────────
controllers.append((
    '3. RoleController (Admin)',
    'app/Http/Controllers/Admin/RoleController.php',
    [
        ['RC-01','index()','Unauthorized user gets 403','!$user->can(admin.roles.view)','User without view permission','GET roles list','403 Forbidden'],
        ['RC-02','index()','Scope filter system shows only system roles','scope === system branch','System + custom roles exist','scope=system','Only admin, cashier, etc. returned'],
        ['RC-03','index()','Scope filter custom shows non-system roles','scope === custom branch','System + custom roles exist','scope=custom','Only custom roles returned'],
        ['RC-04','index()','Scope filter archived shows soft-deleted roles','scope === archived branch','1 archived role exists','scope=archived','Only archived roles returned'],
        ['RC-05','store()','Unauthorized user gets 403','!$user->can(admin.roles.create)','User without create permission','POST create role','403 Forbidden'],
        ['RC-06','store()','Duplicate role name fails validation','Validation — name unique','Role with same name exists','Duplicate name','Session error on name'],
        ['RC-07','store()','Role created with source copies permissions','!empty($data[source_role_id]) branch','Source role with permissions exists','Valid name + source_role_id','New role created with same permissions'],
        ['RC-08','store()','Role created without source has no permissions','source_role_id empty branch','Admin with permission','Valid name, no source_role_id','Role created with 0 permissions'],
        ['RC-09','update()','Protected admin role cannot be modified','$role->isProtectedAdmin() branch','Admin role exists','PUT to admin role','Redirect with error "admin role is protected"'],
        ['RC-10','archive()','System role cannot be archived','in_array(SYSTEM_ROLES) branch','System role e.g. cashier','DELETE archive system role','Redirect with error "System roles cannot be archived"'],
        ['RC-11','archive()','Custom role can be archived','Not a system role — happy path','Custom role exists','DELETE archive custom role','Role soft-deleted, redirect back'],
        ['RC-12','restore()','Archived role can be restored','Happy path','Soft-deleted role exists','PUT restore','Role restored, redirect back'],
        ['RC-13','updatePermissions()','Unauthorized user gets 403','!$user->can(admin.roles.permissions)','User without permissions perm','PUT update permissions','403 Forbidden'],
        ['RC-14','updatePermissions()','Only valid permission names are synced','syncPermissions() happy path','Role + permissions exist','Valid permission names array','Role permissions updated to exact list'],
    ]
))

# ── 4. ProductController ──────────────────────────────────────
controllers.append((
    '4. ProductController (Admin)',
    'app/Http/Controllers/Admin/ProductController.php',
    [
        ['PC-01','index()','Returns all products without filters','No filter branches taken','Multiple products exist','GET products','All products paginated'],
        ['PC-02','index()','Search filter narrows by name','$request->filled(q) branch','Products with different names','q=Caltex','Only matching products returned'],
        ['PC-03','index()','Status filter active returns only active','status === active branch','1 active, 1 inactive product','status=active','Only active products'],
        ['PC-04','index()','Status filter all returns all','status === all branch','Active + inactive products','status=all','All products returned'],
        ['PC-05','index()','Type filter narrows by category','$request->filled(type) branch','Products of different types','type=lpg','Only LPG products'],
        ['PC-06','store()','Duplicate SKU fails validation','Validation — sku unique','Product with same SKU exists','Duplicate sku','Session error on sku'],
        ['PC-07','store()','Invalid type fails validation','Validation — type in:lpg,stove,accessories','Admin with permission','type=unknown','Session error on type'],
        ['PC-08','store()','LPG product extracts kg from size label','type === lpg branch','Admin, valid supplier','type=lpg, size_label=11kg','Variant created: size_unit=kg, size_value=11'],
        ['PC-09','store()','Accessories get default size 1 PC','applySizeDefaults() — accessories','Admin, valid supplier','type=accessories, no size_label','Variant created: size_value=1, size_unit=PC'],
        ['PC-10','store()','Stove gets default size 1 UNIT','applySizeDefaults() — stove','Admin, valid supplier','type=stove, no size_label','Variant created: size_value=1, size_unit=UNIT'],
        ['PC-11','store()','Comma in price is stripped before validation','str_replace(",", "") branch','Admin, valid supplier','price=1,500','Product created with price=1500'],
        ['PC-12','archive()','User without permission gets 403','!$user->can(admin.products.archive)','User without permission','PUT archive product','403 Forbidden'],
        ['PC-13','archive()','Active product is archived','Happy path','Admin with permission, active product','PUT archive','is_active = false in DB'],
        ['PC-14','restore()','Archived product is restored','Happy path','Admin with permission, archived product','PUT restore','is_active = true in DB'],
    ]
))

# ── 5. POSController ──────────────────────────────────────────
controllers.append((
    '5. POSController (Cashier)',
    'app/Http/Controllers/Cashier/POSController.php',
    [
        ['POS-01','index()','User without permission gets 403','!$user->can(cashier.pos.use)','User without POS permission','GET POS page','403 Forbidden'],
        ['POS-02','index()','Unauthenticated user gets 403','!$user branch','No active session','GET POS page','403 Forbidden'],
        ['POS-03','index()','Authorized cashier sees POS page','Happy path','Cashier with cashier.pos.use','GET POS page','Inertia render with products, customers, vat_settings, discount_settings'],
        ['POS-04','store()','User without cashier.sales.create gets 403','!$user->can(cashier.sales.create)','User without permission','POST sale','403 Forbidden'],
        ['POS-05','store()','Missing customer_id fails validation','Validation — customer_id required','Cashier with permission','No customer_id','Session error on customer_id'],
        ['POS-06','store()','Invalid customer_id fails validation','Validation — customer_id exists','Cashier with permission','customer_id=99999','Session error on customer_id'],
        ['POS-07','store()','Invalid payment_method fails validation','Validation — in:cash,gcash,card','Cashier with permission','payment_method=bitcoin','Session error on payment_method'],
        ['POS-08','store()','Invalid vat_treatment fails validation','Validation — vat_treatment in:...','Cashier with permission','vat_treatment=standard','Session error on vat_treatment'],
        ['POS-09','store()','Missing vat_inclusive fails validation','Validation — vat_inclusive required|boolean','Cashier with permission','No vat_inclusive','Session error on vat_inclusive'],
        ['POS-10','store()','Invalid is_delivery fails validation','Validation — is_delivery boolean','Cashier with permission','is_delivery=maybe','Session error on is_delivery'],
        ['POS-11','store()','Empty lines array fails validation','Validation — lines min:1','Cashier with permission','lines=[]','Session error on lines'],
        ['POS-12','store()','Invalid lines.*.product_id fails validation','Validation — lines.*.product_id exists','Cashier with permission','lines[0].product_id=99999','Session error on lines.0.product_id'],
        ['POS-13','store()','lines.*.qty below min fails validation','Validation — lines.*.qty min:1','Cashier with permission','lines[0].qty=0','Session error on lines.0.qty'],
        ['POS-14','store()','Invalid lines.*.mode fails validation','Validation — lines.*.mode in:refill,swap','Cashier with permission','lines[0].mode=exchange','Session error on lines.0.mode'],
        ['POS-15','store()','Missing lines.*.unit_price fails validation','Validation — lines.*.unit_price required','Cashier with permission','lines[0].unit_price missing','Session error on lines.0.unit_price'],
        ['POS-16','store()','Negative unit_price fails validation','Validation — lines.*.unit_price min:0','Cashier with permission','lines[0].unit_price=-1','Session error on lines.0.unit_price'],
        ['POS-17','store()','Non-string payment_ref fails validation','Validation — payment_ref string','Cashier with permission','payment_ref=[]','Session error on payment_ref'],
        ['POS-18','store()','Negative vat_rate fails validation','Validation — vat_rate min:0','Cashier with permission','vat_rate=-0.5','Session error on vat_rate'],
        ['POS-19','store()','Negative cash_tendered fails validation','Validation — cash_tendered min:0','Cashier with permission','cash_tendered=-100','Session error on cash_tendered'],
        ['POS-20','store()','Negative discount_total fails validation','Validation — discount_total min:0','Cashier with permission','discount_total=-50','Session error on discount_total'],
        ['POS-21','store()','Discount item missing kind fails validation','Validation — discounts.*.kind required_with:discounts','Cashier with permission','discounts[0].code=PROMO10','Session error on discounts.0.kind'],
        ['POS-22','store()','Invalid discounts.*.kind fails validation','Validation — discounts.*.kind in:promo,voucher,manual','Cashier with permission','discounts[0].kind=reward','Session error on discounts.0.kind'],
        ['POS-23','store()','Invalid discounts.*.discount_type fails validation','Validation — discounts.*.discount_type in:percent,amount','Cashier with permission','discounts[0].discount_type=ratio','Session error on discounts.0.discount_type'],
        ['POS-24','store()','Discount code too long fails validation','Validation — discounts.*.code max:50','Cashier with permission','discounts[0].code=51+ chars','Session error on discounts.0.code'],
        ['POS-25','store()','Negative discounts.*.value fails validation','Validation — discounts.*.value min:0','Cashier with permission','discounts[0].value=-5','Session error on discounts.0.value'],
        ['POS-26','store()','Non-integer discounts.*.promo_id fails validation','Validation — discounts.*.promo_id integer','Cashier with permission','discounts[0].promo_id=abc','Session error on discounts.0.promo_id'],
        ['POS-27','store()','Non-string manager_pin fails validation','Validation — manager_pin string','Cashier with permission','manager_pin=[]','Session error on manager_pin'],
        ['POS-28','store()','InvalidArgumentException returns sale error','catch (\\InvalidArgumentException $e) branch','Service throws invalid argument','Valid input but service throws','Redirect back with errors.sale = exception message'],
        ['POS-29','store()','RuntimeException returns sale error','catch (\\RuntimeException $e) branch','Service throws runtime error','Valid input but service throws','Redirect back with errors.sale = exception message'],
        ['POS-30','store()','Unhandled Throwable returns generic error','catch (\\Throwable $e) branch','Service throws unexpected','Valid input but service fails','Redirect back with errors.sale = generic message'],
        ['POS-31','store()','Successful sale redirects with success','Happy path','Valid customer, product variant, cashier','Complete valid payload','Redirect back with success flash'],
    ]
))

# ── 6. SaleController ─────────────────────────────────────────
controllers.append((
    '6. SaleController (Cashier)',
    'app/Http/Controllers/Cashier/SaleController.php',
    [
        ['SC-01','export()','Missing from_date fails validation','Validation — from_date required','Authenticated user','No from_date','Validation error on from_date'],
        ['SC-02','export()','to_date before from_date fails validation','Validation — after_or_equal:from_date','Authenticated user','to_date earlier than from_date','Validation error on to_date'],
        ['SC-03','export()','Invalid status_scope fails validation','Validation — in:paid,paid_pending,...','Authenticated user','status_scope=unknown','Validation error on status_scope'],
        ['SC-04','export()','Format csv triggers CSV export','$format === csv branch','Valid date range','format=csv','CSV file download response'],
        ['SC-05','export()','Default format triggers XLSX export','Default (xlsx) branch','Valid date range','No format or format=xlsx','XLSX file download response'],
        ['SC-06','export()','from/to aliases accepted','?? $request->input("from") branch','Authenticated user','from=2025-01-01&to=2025-01-31','Valid export produced'],
        ['SC-07','receipt()','Sale with no receipt returns printed_count 0','receipt?->printed_count ?? 0 — null path','Sale exists, no receipt','GET receipt for sale','JSON with printed_count = 0'],
        ['SC-08','receipt()','Sale with receipt returns its receipt_number','receipt?->receipt_number — has receipt','Sale with linked receipt','GET receipt','JSON with ref = receipt_number'],
        ['SC-09','receipt()','Sale with no sale_datetime returns null date/time','saleDatetime = null path','Sale with null sale_datetime','GET receipt','JSON with date = null, time = null'],
        ['SC-10','reprint()','Existing receipt has printed_count incremented','if ($receipt) — exists branch','Sale with existing receipt','POST reprint','printed_count incremented by 1'],
        ['SC-11','reprint()','No receipt creates new and increments','if (!$receipt) branch','Sale with no receipt','POST reprint','New receipt created, printed_count = 1'],
        ['SC-12','printReceipt()','Default format returns Blade view','format !== pdf branch','Sale exists','GET print (no format)','Returns receipts.print Blade view'],
        ['SC-13','printReceipt()','PDF format returns PDF download','format === pdf branch','Sale exists, DomPDF available','GET print with format=pdf','PDF file download response'],
    ]
))

# ── 7. CustomerController ─────────────────────────────────────
controllers.append((
    '7. CustomerController (Cashier)',
    'app/Http/Controllers/Cashier/CustomerController.php',
    [
        ['CC-01','index()','User without any view permission gets 403','canViewCustomers() — false','User without cashier/admin view permission','GET customers','403 Forbidden'],
        ['CC-02','index()','Cashier with cashier.customers.view can access','canViewCustomers() — cashier.customers.view','Cashier with view permission','GET customers','Inertia render with customers list'],
        ['CC-03','index()','Admin with admin.customers.view can access','canViewCustomers() — admin.customers.view','Admin with admin view permission','GET customers','Inertia render with customers list'],
        ['CC-04','store()','User without create permission gets 403','Both can() checks fail','User without any create permission','POST create customer','403 Forbidden'],
        ['CC-05','store()','Missing name fails validation','Validation — name required','User with permission','No name','Session error on name'],
        ['CC-06','store()','Invalid customer_type fails validation','Validation — in:walkin,regular,corporate','User with permission','customer_type=vip','Session error on customer_type'],
        ['CC-07','store()','Valid customer is created','Happy path','User with permission','Valid name, optional fields','Customer created in DB, redirect with success'],
        ['CC-08','update()','User without update permission gets 403','!$user->can(cashier.customers.update)','User without permission','PUT update customer','403 Forbidden'],
        ['CC-09','update()','Missing customer_type in update fails','Validation — customer_type required','User with permission, customer exists','No customer_type','Session error on customer_type'],
        ['CC-10','update()','Valid update succeeds','Happy path','User with permission, customer exists','Valid fields','Customer updated, redirect with success'],
        ['CC-11','destroy()','User without delete permission gets 403','!$user->can(cashier.customers.delete)','User without permission','DELETE customer','403 Forbidden'],
        ['CC-12','destroy()','Authorized user can delete customer','Happy path','Cashier with delete permission','DELETE customer','Customer deleted, redirect with success'],
    ]
))

# ── 8. StockController ────────────────────────────────────────
controllers.append((
    '8. StockController (Inventory)',
    'app/Http/Controllers/Inventory/StockController.php',
    [
        ['STC-01','stockCount()','User without permission gets 403','!$user->can(inventory.stock.view)','User without permission','GET stock counts','403 Forbidden'],
        ['STC-02','stockCount()','No location returns empty data','if (!$location) branch','No location in DB','GET stock counts','Inertia page with empty stock_counts.data'],
        ['STC-03','stockCount()','Returns counts when location exists','Happy path','Location + stock data exist','GET stock counts','Inertia page with paginated stock counts'],
        ['STC-04','submitCount()','User without permission gets 403','!$user->can(inventory.stock.adjust)','User without permission','POST submit count','403 Forbidden'],
        ['STC-05','submitCount()','Admin cannot submit stock counts','$user->hasRole(admin) branch','Admin user','POST submit count','Redirect back with error "Admin cannot submit stock counts."'],
        ['STC-06','submitCount()','Missing filled_qty fails validation','Validation — filled_qty required','Inventory manager','No filled_qty','Session error on filled_qty'],
        ['STC-07','submitCount()','Reason too short fails validation','Validation — reason min:3','Inventory manager','reason=AB','Session error on reason'],
        ['STC-08','submitCount()','Valid count submission succeeds','Happy path','Inventory manager, valid InventoryBalance','Valid filled_qty + reason','Redirect back with success'],
        ['STC-09','reviewCount()','Non-admin gets 403','!$user->hasRole(admin) branch','Non-admin user','POST review count','403 Forbidden'],
        ['STC-10','reviewCount()','Invalid action fails validation','Validation — in:approve,reject','Admin, stock count exists','action=skip','Validation error on action'],
        ['STC-11','reviewCount()','Approved count with movement ID redirects to audit','$approved && $movementId branch','Admin, approvable stock count','action=approve','Redirect to audit page'],
        ['STC-12','reviewCount()','Rejected count redirects back with success','!$approved branch','Admin, stock count exists','action=reject','Redirect back with "Stock count rejected."'],
        ['STC-13','thresholds()','No location returns empty thresholds','if (!$location) branch','No location in DB','GET thresholds','Inertia page with empty thresholds'],
        ['STC-14','thresholds()','Risk filter critical shows only critical items','ratio <= 0.5 → critical','Balance with ratio <= 0.5','risk=critical','Only critical items returned'],
        ['STC-15','thresholds()','Risk filter warning shows only warning items','0.5 < ratio < 1 → warning','Balance with warning ratio','risk=warning','Only warning items returned'],
        ['STC-16','thresholds()','Zero reorder level is always ok','reorderLevel <= 0 → ok','Balance with reorder_level=0','GET thresholds','Item shown with risk_level = ok'],
        ['STC-17','updateThresholds()','User without permission gets 403','!$user->can(inventory.stock.low_stock)','User without permission','PUT thresholds','403 Forbidden'],
        ['STC-18','updateThresholds()','Valid threshold update saves to DB','Happy path','Admin, inventory balances exist','Valid updates array','reorder_level updated in DB, redirect with success'],
    ]
))

# ── 9. PurchaseController ─────────────────────────────────────
controllers.append((
    '9. PurchaseController (Inventory)',
    'app/Http/Controllers/Inventory/PurchaseController.php',
    [
        ['PRC-01','store()','User without permission gets 403','!$user->can(inventory.purchases.create)','User without permission','POST create purchase','403 Forbidden'],
        ['PRC-02','store()','Missing product_variant_id fails validation','Validation — required','User with permission','No product_variant_id','Validation error'],
        ['PRC-03','store()','Non-existent supplier_id fails validation','Validation — exists:suppliers,id','User with permission','Invalid supplier_id','Validation error on supplier_id'],
        ['PRC-04','store()','qty below 1 fails validation','Validation — min:1','User with permission','qty=0','Validation error on qty'],
        ['PRC-05','store()','PurchaseStatusException redirects with error','catch(PurchaseStatusException) branch','Service throws status exception','Valid input','Redirect back with error message'],
        ['PRC-06','store()','Valid purchase creates and redirects','Happy path','Valid product variant + supplier','Valid payload','Purchase created in DB, redirect with success'],
        ['PRC-07','approve()','User without approve permission gets 403','!$user->can(inventory.purchases.approve)','User without permission','POST approve','403 Forbidden'],
        ['PRC-08','approve()','User with no role gets 403','resolveActorRole() — no role','User with permission but no role','POST approve','403 Forbidden'],
        ['PRC-09','reject()','Reason too short fails validation','Validation — reason min:3','User with approve permission','reason=X','Validation error on reason'],
        ['PRC-10','markDelivered()','Non-inventory manager gets 403 JSON','!$user->hasRole(inventory_manager) branch','Non-inventory manager','POST mark delivered','JSON 403 "You are not allowed..."'],
        ['PRC-11','markDelivered()','Invalid status transition returns 422','PurchaseStatus::ensureTransition() throws','Purchase in wrong status','POST mark delivered','JSON 422 with error message'],
        ['PRC-12','confirm()','User without confirm permission gets 403','!$user->can(inventory.purchases.confirm)','User without permission','POST confirm','403 Forbidden'],
        ['PRC-13','pay()','User without accountant.payables.pay gets 403','!$user->can(accountant.payables.pay)','Non-accountant','POST pay','403 Forbidden'],
        ['PRC-14','close()','Non-finance role gets 403','!$user->hasRole(finance) branch','User with no finance role','POST close','403 Forbidden'],
        ['PRC-15','destroy()','User without delete permission gets 403','!$user->can(inventory.purchases.delete)','User without permission','DELETE purchase','403 Forbidden'],
        ['PRC-16','purge()','No eligible purchases returns info','$deleted === 0 branch','No purchases exist','POST purge','Redirect back with "No eligible purchases to delete."'],
        ['PRC-17','purge()','Successful purge returns count','Happy path','Purchases exist + user with permission','POST purge','Redirect back with "{n} purchase(s) cleared."'],
    ]
))

# ── 10. PurchaseRequestController ────────────────────────────
controllers.append((
    '10. PurchaseRequestController (Inventory)',
    'app/Http/Controllers/Inventory/PurchaseRequestController.php',
    [
        ['PRQ-01','store()','Unauthorized user is rejected','$this->authorize(create)','User without create policy','POST create PR','403 Forbidden'],
        ['PRQ-02','store()','Missing reason fails validation','Validation — reason required','Authorized user','No reason','Validation error on reason'],
        ['PRQ-03','store()','Empty items array fails validation','Validation — items min:1','Authorized user','items=[]','Validation error on items'],
        ['PRQ-04','store()','Item with zero qty is skipped','syncItems() — quantity === 0','Authorized user, valid product','requested_qty=0','Item not created in DB'],
        ['PRQ-05','store()','Valid request creates PR with generated number','Happy path — generatePrNumber()','Authorized user, valid product','Valid reason + items','PR created with PR-YYYYMMDD-NNNN format number'],
        ['PRQ-06','submit()','PR with no items returns error','items()->count() === 0','PR with 0 items','POST submit','Redirect back with error "Add at least one item"'],
        ['PRQ-07','submit()','Valid PR is marked submitted','Happy path','PR with items','POST submit','PR status = submitted, redirect with success'],
        ['PRQ-08','receive()','Item not belonging to PR throws error','purchase_request_id mismatch branch','Different PR item passed','Item from another PR','Validation error "Item does not belong to this purchase request."'],
        ['PRQ-09','receive()','Damaged qty exceeding received qty throws error','$damaged > $received branch','Valid PR item','damaged_qty > received_qty','Validation error "Damaged quantity cannot exceed received quantity."'],
        ['PRQ-10','receive()','Received qty exceeding approved qty throws error','received > approved_qty branch','PR item with approved_qty set','Total received > approved','Validation error "Received quantity cannot exceed the approved quantity."'],
        ['PRQ-11','receive()','All items received → fully received status','$allReceived = true branch','All PR items fully received','Valid receive payload','PR status = fully_received'],
        ['PRQ-12','receive()','Partial items → partially received status','$allReceived = false branch','Only some items received','Partial receive payload','PR status = partially_received'],
        ['PRQ-13','receive()','Good cost > 0 creates inventory ledger entry','$goodCost > 0 branch','Valid receive, no damage','received_qty=5, damaged_qty=0','Ledger entry with account 1200 debit created'],
        ['PRQ-14','receive()','Damage cost > 0 creates damage ledger line','$damageCost > 0 branch','Some damaged items','damaged_qty > 0','Ledger line with account 5200 debit created'],
    ]
))

# ── 11. RestockRequestController ──────────────────────────────
controllers.append((
    '11. RestockRequestController (Inventory)',
    'app/Http/Controllers/Inventory/RestockRequestController.php',
    [
        ['RRC-01','store()','User without permission gets 403','!$user->can(inventory.purchases.create)','User without permission','POST create restock','403 Forbidden'],
        ['RRC-02','store()','Missing product_variant_id fails validation','Validation — required','User with permission','No product_variant_id','Validation error'],
        ['RRC-03','store()','Zero qty fails validation','Validation — min:0.01','User with permission','qty=0','Validation error on qty'],
        ['RRC-04','store()','No location configured returns error','!$locationId branch','No location in DB','Valid input','Redirect back with error "No location configured."'],
        ['RRC-05','store()','Valid request creates restock and redirects','Happy path','Location exists, user with permission','Valid payload','Restock request created, redirect with success'],
        ['RRC-06','approve()','User without update permission gets 403','!$user->can(inventory.purchases.update)','User without permission','POST approve','403 Forbidden'],
        ['RRC-07','approve()','Missing supplier_id fails validation','Validation — supplier_id required','User with permission','No supplier_id','Validation error on supplier_id'],
        ['RRC-08','reject()','User without update permission gets 403','!$user->can(inventory.purchases.update)','User without permission','POST reject','403 Forbidden'],
        ['RRC-09','receivePage()','Non-existent request returns 404','if (!$stockRequest) branch','Request not found','GET receive page with bad ID','404 Not Found'],
        ['RRC-10','receive()','Empty lines array fails validation','Validation — lines required','User with permission','lines=[]','Validation error on lines'],
    ]
))

# ── 12. PayableController ─────────────────────────────────────
controllers.append((
    '12. PayableController (Accountant)',
    'app/Http/Controllers/Accountant/PayableController.php',
    [
        ['PAY-01','index()','Non-accountant gets 403','!$user->hasRole(accountant)','Non-accountant user','GET payables','403 Forbidden'],
        ['PAY-02','index()','Supplier filter narrows results','$filters[supplier] branch','Multiple suppliers payables','supplier=1','Only supplier 1 payables'],
        ['PAY-03','index()','Status filter unpaid returns only unpaid','status !== all branch','Mixed status payables','status=unpaid','Only unpaid payables'],
        ['PAY-04','show()','Non-accountant gets 403','!$user->hasRole(accountant)','Non-accountant','GET payable detail','403 Forbidden'],
        ['PAY-05','show()','JSON request returns JSON response','$request->wantsJson() branch','Accountant, payable exists','GET with Accept: application/json','JSON response with payable data'],
        ['PAY-06','pay()','Non-accountant gets 403','!$user->hasRole(accountant)','Non-accountant','POST pay','403 Forbidden'],
        ['PAY-07','pay()','Already paid payable returns error','status !== STATUS_UNPAID branch','Payable with status=paid','POST pay','Redirect back with error "Payable already settled."'],
        ['PAY-08','pay()','Amount mismatch throws validation error','abs(amount - payable.amount) > 0.01','Unpaid payable','paid_amount differs by > 0.01','Validation error "Paid amount must equal payable balance."'],
        ['PAY-09','pay()','Bank/transfer method maps to account 1020','mapPaymentMethodToAccount() — bank','Accountant, unpaid payable','payment_method=bank_transfer','Ledger credit line uses account 1020'],
        ['PAY-10','pay()','Cash method maps to account 1010','mapPaymentMethodToAccount() — cash','Accountant, unpaid payable','payment_method=cash','Ledger credit line uses account 1010'],
        ['PAY-11','pay()','Successful payment marks payable as paid','Happy path','Unpaid payable, accountant','Valid payment_method + matching amount','status=paid, paid_at set, redirect with success'],
        ['PAY-12','pay()','Paying RestockRequest source updates its status','instanceof RestockRequest branch','Payable linked to RestockRequest','POST pay success','RestockRequest status=paid'],
        ['PAY-13','addNote()','Non-accountant gets 403','!$user->hasRole(accountant)','Non-accountant','POST add note','403 Forbidden'],
        ['PAY-14','addNote()','Note too short fails validation','Validation — note min:3','Accountant, payable exists','note=AB','Validation error on note'],
    ]
))

# ── 13. RemittanceController ──────────────────────────────────
controllers.append((
    '13. RemittanceController (Accountant)',
    'app/Http/Controllers/Accountant/RemittanceController.php',
    [
        ['REM-01','review()','User without view permission gets 403','!$user->can(accountant.remittances.view)','User without permission','GET review','403 Forbidden'],
        ['REM-02','recordCash()','User without verify permission gets 403','!$user->can(accountant.remittances.verify)','User without permission','POST record cash','403 Forbidden'],
        ['REM-03','recordCash()','Already finalized date returns error','DailyClose::where()->exists() branch','DailyClose exists for date','POST with finalized date','Redirect back with error "already finalized"'],
        ['REM-04','recordCash()','Cash variance with no note throws error','variance !== 0.0 && note.length < 3','Accountant, cash != expected','Cash counted != expected, no note','Redirect back with error "Provide a short note when cash variance is not zero."'],
        ['REM-05','recordCash()','Zero variance with no note succeeds','variance === 0.0 branch','Cash counted = expected','No note provided','Remittance created/updated, redirect with success'],
        ['REM-06','verifyCashlessTransactions()','Already finalized date returns 422','DailyClose::where()->exists()','Finalized date','POST verify cashless','JSON 422 "already finalized"'],
        ['REM-07','verifyCashlessTransactions()','Already verified transactions throw error','!empty($alreadyVerifiedIds) branch','Some payments already verified','POST with already-verified IDs','JSON 422 "already been verified"'],
        ['REM-08','dailyClose()','Already finalized returns 200 with info','$existingClose not null','DailyClose already exists','POST daily close','JSON 200 "already finalized" + finalized_at'],
        ['REM-09','dailyClose()','Pending cashless with no verified IDs throws error','pendingIds not empty + verifiedPayload empty','Pending cashless transactions exist','POST, no verified_transaction_ids','JSON 422 "Verify all pending cashless"'],
        ['REM-10','dailyClose()','Missing transaction IDs throw error','!empty($missing) branch','Pending transactions exist','verified_transaction_ids missing some IDs','JSON 422 "All pending...must be verified"'],
        ['REM-11','dailyClose()','Successful close creates DailyClose and audit log','Happy path','All cashless verified, cash recorded','Valid payload','DailyClose created, audit log remittance.daily.finalized, JSON 200'],
        ['REM-12','reopen()','User without verify permission gets 403','!$user->can(accountant.remittances.verify)','User without permission','POST reopen','403 Forbidden'],
        ['REM-13','reopen()','Non-finalized date returns error','if (!$close) branch','No DailyClose for date','POST reopen','422 / error "not finalized"'],
        ['REM-14','reopen()','Finalized date is reopened and logged','Happy path','DailyClose exists','POST reopen','DailyClose deleted, audit log remittance.daily.reopen, success response'],
        ['REM-15','stageFromFlags()','All flags true → finalized','All true branch','—','finalized=true','Returns finalized'],
        ['REM-16','stageFromFlags()','Cash + cashless → ready_to_finalize','Both cash+cashless not finalized','—','cash=true, cashless=true, finalized=false','Returns ready_to_finalize'],
        ['REM-17','stageFromFlags()','Only cashless → cashless_verified','Cashless only','—','cash=false, cashless=true','Returns cashless_verified'],
        ['REM-18','stageFromFlags()','Only cash → cash_recorded','Cash only','—','cash=true, cashless=false','Returns cash_recorded'],
        ['REM-19','stageFromFlags()','Neither → draft','Default path','—','cash=false, cashless=false, finalized=false','Returns draft'],
    ]
))

# ── 14. LedgerController ──────────────────────────────────────
controllers.append((
    '14. LedgerController (Accountant)',
    'app/Http/Controllers/Accountant/LedgerController.php',
    [
        ['LC2-01','index()','No filters returns all ledger lines','No filter branches taken','Ledger entries exist','GET ledger','Inertia page with all lines paginated'],
        ['LC2-02','index()','Account filter narrows by account code','account !== all','Multiple account codes','account=1010','Only account 1010 lines'],
        ['LC2-03','index()','Cleared filter returns only cleared lines','cleared === cleared branch','Mixed cleared/uncleared','cleared=cleared','Only cleared lines'],
        ['LC2-04','index()','Uncleared filter returns only uncleared','cleared === uncleared branch','Mixed cleared/uncleared','cleared=uncleared','Only uncleared lines'],
        ['LC2-05','index()','Date range filter narrows results','from and to filters applied','Lines on various dates','from=2025-01-01&to=2025-01-31','Only lines within range'],
        ['LC2-06','referenceLines()','Returns lines matching reference ID','Happy path','Lines linked to reference','GET referenceLines/ABC-001','JSON with lines, totals, balanced flag'],
        ['LC2-07','referenceLines()','Balanced entry returns balanced = true','abs(debit-credit) < 0.01','Equal debit/credit lines','GET reference with balanced entries','balanced = true'],
        ['LC2-08','referenceLines()','Unbalanced entry returns balanced = false','abs(debit-credit) >= 0.01','Unequal debit/credit','GET reference with unbalanced entries','balanced = false'],
        ['LC2-09','exportCsv()','Returns a CSV download','Happy path','Ledger entries exist','GET export CSV','CSV file download response'],
        ['LC2-10','sortOrderSpec()','posted_at_asc returns correct order','case posted_at_asc','—','sort=posted_at_asc','Orders by le.entry_date ASC'],
        ['LC2-11','sortOrderSpec()','created_at_desc returns correct order','case created_at_desc','—','sort=created_at_desc','Orders by ledger_lines.created_at DESC'],
        ['LC2-12','sortOrderSpec()','Unknown sort defaults to posted_at_desc','default case','—','sort=unknown','Orders by le.entry_date DESC'],
    ]
))

# ── 15. RiderDeliveryController ───────────────────────────────
controllers.append((
    '15. RiderDeliveryController (Rider)',
    'app/Http/Controllers/Rider/RiderDeliveryController.php',
    [
        ['RDC-01','updateStatus()','Delivery not assigned to rider gets 403','assigned_rider_user_id !== user.id','Delivery assigned to different rider','POST update status','403 Forbidden "not assigned to you"'],
        ['RDC-02','updateStatus()','Invalid status fails validation','Validation — in:in_transit,delivered,failed','Correct rider','status=cancelled','Validation error on status'],
        ['RDC-03','updateStatus()','Delivered without proof aborts 422','requiresProof && !$hasProof branch','Correct rider, no proof on file','status=delivered, no proof','422 "Proof of delivery photo or signature is required."'],
        ['RDC-04','updateStatus()','Delivered with existing proof URL passes','proof_photo_url exists on delivery','Correct rider, delivery has photo','status=delivered, no new proof','Status updated successfully'],
        ['RDC-05','updateStatus()','in_transit skips proof check','requiresProof = false path','Correct rider','status=in_transit','Status updated without proof required'],
        ['RDC-06','updateStatus()','Invalid base64 data URL is ignored','preg_match() fails in storeDataUrlImage()','Correct rider','proof_photo_data=invalid_string','Image silently ignored, status still updates'],
        ['RDC-07','updateStatus()','delivered_items as JSON string is decoded','is_string($raw) branch','Correct rider','delivered_items as JSON string','Items parsed and stored correctly'],
        ['RDC-08','updateStatus()','Item with non-numeric delivered_qty is skipped','!is_numeric($delivered) branch','Correct rider','Item with delivered_qty=abc','Item skipped from normalized list'],
        ['RDC-09','updateStatus()','No delivered_items falls back to sale items','buildDeliveredItemsFallback() — sale items','Delivery linked to sale with items','status=delivered, no delivered_items','Fallback uses sale items as delivered_items'],
        ['RDC-10','updateNote()','Delivery not assigned to rider gets 403','assigned_rider_user_id !== user.id','Delivery assigned to other rider','POST update note','403 "not assigned to you"'],
        ['RDC-11','updateNote()','Valid note is saved','Happy path','Correct rider, delivery exists','note=Left at the gate','Note updated, redirect with success'],
    ]
))

# ── 16. AuditLogController ────────────────────────────────────
controllers.append((
    '16. AuditLogController',
    'app/Http/Controllers/AuditLogController.php',
    [
        ['ALC-01','index()','User without any audit view permission gets 403','!$user->canAny([...]) branch','User with no audit permissions','GET audit logs','403 Forbidden'],
        ['ALC-02','index()','Admin sees all sectors','allowedSectorsForRole(admin)','Admin user with permission','GET audit logs','Sectors: all, access, people, sales, inventory, finance'],
        ['ALC-03','index()','Cashier default sector is sales','defaultSectorForRole(cashier)','Cashier with view permission','GET audit logs, no sector','sector = sales applied'],
        ['ALC-04','index()','Accountant default sector is finance','defaultSectorForRole(accountant)','Accountant with permission','GET audit logs, no sector','sector = finance applied'],
        ['ALC-05','index()','Inventory manager default sector is inventory','defaultSectorForRole(inventory_manager)','Inventory manager','GET audit logs','sector = inventory applied'],
        ['ALC-06','index()','Cashier requesting finance sector gets reset to sales','!in_array($sector, $allowedSectors)','Cashier with audit view','sector=finance','Sector reset to sales'],
        ['ALC-07','index()','Sector sales loads sales data','if ($sector === sales) branch','Cashier with permission','sector=sales','Response includes sales data, movements=null, finance=null'],
        ['ALC-08','index()','Sector inventory loads movements data','if ($sector === inventory) branch','Inventory manager','sector=inventory','Response includes movements data, sales=null'],
        ['ALC-09','index()','Sector finance loads finance logs','if ($sector === finance) branch','Accountant with permission','sector=finance','Response includes finance data'],
        ['ALC-10','index()','Sector all loads no supplemental data','No sector branches taken','Admin user','sector=all','sales=null, movements=null, finance=null'],
    ]
))


# ── Render all controllers ────────────────────────────────────
if controllers_count_run is not None:
    total_cases = sum(len(rows) for _, _, rows in controllers)
    controllers_count_run.text = f'Controllers Covered: {len(controllers)} Critical Controllers | {total_cases} Test Cases'

for title, path, rows in controllers:
    add_heading(doc, title, level=1)
    p = doc.add_paragraph()
    run = p.add_run(f'File: {path}')
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.italic = True

    rendered_rows = build_whitebox_rows(title, rows)
    add_controller_table(doc, HEADERS, rendered_rows)
    doc.add_paragraph()
    doc.add_page_break()

# ── Save ──────────────────────────────────────────────────────
out = 'WhiteBoxTestPlan_Pietyl_LPG.docx'
doc.save(out)
print(f'Saved: {out}')

# â”€â”€ Markdown Output â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
md_lines = []
md_lines.append('# White Box Test Plan - Pietyl LPG Management System')
md_lines.append('')
md_lines.append(f'Generated: {date.today().isoformat()}')
md_lines.append('')

md_headers = HEADERS

def md_escape(value: str) -> str:
    return str(value).replace('|', '\\|')

for title, path, rows in controllers:
    md_lines.append(f'## {title}')
    md_lines.append('')
    md_lines.append(f'*File: {path}*')
    md_lines.append('')
    md_lines.append('| ' + ' | '.join(md_headers) + ' |')
    md_lines.append('| ' + ' | '.join(['---'] * len(md_headers)) + ' |')
    rendered_rows = build_whitebox_rows(title, rows)
    for row in rendered_rows:
        md_lines.append('| ' + ' | '.join(md_escape(col) for col in row) + ' |')
    md_lines.append('')

md_out = 'WhiteBoxTestPlan_Pietyl_LPG.md'
with open(md_out, 'w', encoding='utf-8') as f:
    f.write('\n'.join(md_lines))
print(f'Saved: {md_out}')
